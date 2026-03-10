import os
import streamlit as st
from docxtpl import DocxTemplate
import google.generativeai as genai
import json
import re
import time
from dotenv import load_dotenv
import base64
import requests
import smtplib
from email.mime.text import MIMEText
from pathlib import Path

# ==============================================================
# CONFIGURACIÓN INICIAL
# ==============================================================

st.set_page_config(page_title="Generador de Actas", page_icon="📝", layout="wide")

load_dotenv()

# Prioridad: secrets de Streamlit Cloud > variables de entorno
API_KEY = st.secrets.get("GOOGLE_API_KEY") or os.getenv("GOOGLE_API_KEY")
JSONBIN_API_KEY = st.secrets.get("JSONBIN_API_KEY") or os.getenv("JSONBIN_API_KEY")
JSONBIN_BIN_ID = st.secrets.get("JSONBIN_BIN_ID") or os.getenv("JSONBIN_BIN_ID")

if not API_KEY:
    st.error("No se encontró GOOGLE_API_KEY en el archivo .env o en los secretos de Streamlit.")
    st.stop()

genai.configure(api_key=API_KEY)
model = genai.GenerativeModel('gemini-2.5-flash')

TEMPLATES_DIR = Path("templates")
LIMITE_CONTADOR = 13   # <-- límite máximo antes de enviar alerta

# ==============================================================
# CONFIGURACIÓN JSONBIN
# ==============================================================

BASE_URL = f"https://api.jsonbin.io/v3/b/{JSONBIN_BIN_ID}"
HEADERS = {
    "X-Master-Key": JSONBIN_API_KEY or "",
    "Content-Type": "application/json"
}

@st.cache_data(ttl=60)  # Cachea durante 60 segundos para no sobrecargar JSONBin
def obtener_contador():
    try:
        response = requests.get(f"{BASE_URL}/latest", headers=HEADERS, timeout=5)
        response.raise_for_status()
        record = response.json().get("record", {})
        return record.get("contador_actas", 0)
    except Exception as e:
        st.warning(f"⚠️ No se pudo obtener el contador global: {e}")
        return 0

def actualizar_contador(nuevo_valor):
    try:
        response = requests.put(BASE_URL, headers=HEADERS, json={"contador_actas": nuevo_valor}, timeout=5)
        response.raise_for_status()
        st.cache_data.clear()  # Limpia la caché para que el nuevo valor se refleje
    except Exception as e:
        st.error(f"⚠️ No se pudo guardar el contador en JSONBin: {e}")

# ==============================================================
# ALERTA POR CORREO (comentada pero mejorada)
# ==============================================================

# def enviar_alerta_correo(mensaje):
#     user = st.secrets.get("EMAIL_USER") or os.getenv("EMAIL_USER")
#     password = st.secrets.get("EMAIL_PASS") or os.getenv("EMAIL_PASS")
#     destino = st.secrets.get("DESTINO_ALERTA") or os.getenv("DESTINO_ALERTA")
#
#     if not all([user, password, destino]):
#         st.warning("⚠️ No se configuró correctamente el envío de correo (revisa .env o secretos).")
#         return
#
#     msg = MIMEText(mensaje)
#     msg["Subject"] = "⚠️ Alerta: Límite de ACTAS alcanzado"
#     msg["From"] = user
#     msg["To"] = destino
#
#     try:
#         with smtplib.SMTP_SSL("smtp.gmail.com", 465) as server:
#             server.login(user, password)
#             server.send_message(msg)
#         st.info("📨 Se envió una alerta por correo.")
#     except Exception as e:
#         st.error(f"Error al enviar correo: {e}")

# ==============================================================
# CSS PERSONALIZADO
# ==============================================================

st.markdown("""
    <style>
        .app-header {
            display: flex;
            align-items: center;
            justify-content: center;
            gap: 15px;
            margin-bottom: 25px;
            background-color: #ffffff;
            padding: 15px 25px;
            border-radius: 15px;
            box-shadow: 0 2px 8px rgba(0, 0, 0, 0.08);
        }
        .app-header img {
            height: 80px;
            width: auto;
            border-radius: 10px;
        }
        .app-header h1 {
            font-size: 2.2em;
            font-weight: 700;
            color: #1E3A8A;
            margin: 0;
        }
        .section-title {
            font-size: 1.2em;
            font-weight: bold;
            color: #1E40AF;
            margin-top: 25px;
        }
        .footer {
            text-align: center;
            color: #6B7280;
            font-size: 0.9em;
            margin-top: 50px;
            padding-top: 10px;
            border-top: 1px solid #E5E7EB;
        }
        .stButton button {
            background-color: #2563EB;
            color: white;
            border-radius: 8px;
            font-weight: bold;
            padding: 8px 20px;
            transition: all 0.3s ease;
        }
        .stButton button:hover {
            background-color: #1E40AF;
            transform: scale(1.02);
        }
    </style>
""", unsafe_allow_html=True)

# ==============================================================
# FUNCIONES AUXILIARES
# ==============================================================

@st.cache_data
def get_fields_from_template(template_path):
    import docx
    doc = docx.Document(template_path)
    found_fields = set()
    pattern = re.compile(r'\{\{.*?\}\}|\{%.*?%\}')
    for para in doc.paragraphs:
        found_fields.update(pattern.findall(para.text))
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    found_fields.update(pattern.findall(para.text))
    return list(found_fields)

def normalizar_listas(data):
    """Asegura que todos los elementos de las listas tengan los campos requeridos, sin modificar las claves."""
    claves = {
        "ASISTENTES_REUNION": ["nombreasistentereu", "cargoasistentereunion"],
        "TEMAS_TRATADOS_N": ["tema", "responsablet"],
        "COMPROMISOS_DE_REUNION": ["compromiso", "resposablen", "fechac", "fechas"],
        "TEMAS_TRATADOS": ["tema", "desarrollo"],
        "COMPROMISOS_R": ["compromiso", "responsable", "fechaejecucion"],
    }
    for clave, campos in claves.items():
        lista = data.get(clave, [])
        if not isinstance(lista, list):
            lista = []
        for item in lista:
            for campo in campos:
                item.setdefault(campo, "N/A")
        data[clave] = lista  # Mantenemos la clave original en mayúsculas
    return data

@st.cache_data(show_spinner=False)
def extract_info_with_gemini(text_to_process, fields):
    prompt = f"""
    Analiza el siguiente texto y extrae la información para los siguientes campos. 
    Devuelve SOLO un JSON válido. Si no hay información, usa "N/A" o [] según corresponda.

    Campos esperados:
    {', '.join(fields)}

    Instrucciones especificas para campos ESPECIFICOS:
    -   {{FECHA}}: extrae la fecha de la reunión en formato DD/MM/AAAA.
    -	{{HORA_INICIO}}: extrae la hora de inicio de la reunión en formato H:mm.
    -	{{HORA_FIN}}: extrae la hora de fin de la reunión en formato H:mm.
    -	{{CIUDAD}}: extrae la ciudad donde se llevó a cabo la reunión o evento.
    -	{{SEDE}}: extrae la sede (oficina, edificio, empresa, etc.) donde se realizó la reunión.
    -   {{LUGAR_REUNION}}: extrae el lugar donde se realizo la reunion.
    -	{{OBJETIVO_DE_LA_REUNION}}: extrae el objetivo de la reunión explicado de forma clara y completa.
    - {{TEMAS_TRATADOS}}: Esta debe ser una LISTA de objetos JSON. Cada objeto representa un tema tratado en la reunión.
    - Cada objeto debe tener las claves:
        - tema: extrae el tema tratado.
        - desarrollo: extrae de manera detallada como se desarrollo el tema a tratar.
    - {{COMPROMISOS_R}}: Esta debe ser una LISTA de objetos JSON. Cada objeto representa un compromiso de la reunion.
    - Cada objeto debe tener las claves:
        - compromiso: extrae el compromiso a realizar.
        - responsable: extrae el nombre de la persona encargada de ejecutar el compromiso.
        - fechaejecucion: extrae la fecha en la cual se va a ejecutar el compromiso.
    -    {{TEMA_PROXIMA_REUNION}}: extrae el tema a tratar en la proxima reunion.
    -    {{FECHA_PROXIMA_REUNION}}: extrae la fecha en la cual se va a realizar la proxima reunion.
    - {{ASISTENTES_REUNION}}: Esta debe ser una LISTA de objetos JSON. Cada objeto representa una perona que asistio a la reunion.
    - Cada objeto debe tener las claves:
        - nombreasistentereu: extrar el nombre completo de la personas asitente a la reunion.
        - cargoasistentereunion: extrea el cargo de la persona asistente a la reunion.
    - {{TEMAS_TRATADOS_N}}: Esta debe ser una LISTA de objetos JSON. Cada objeto representa un tema tratado en la reunión.
    - Cada objeto debe tener las claves:
        - tema: extrae el tema tratado, haz que los temas relacionados los en listas en uno solo.
        - responsablet: extrae el nombre completo de la persona encargada del tema a tratar.
    - {{DESARROLLO_DE_LA_REUNION_Y_CONCLUSIONES}}: A partir de los temas extraídos en TEMAS_TRATADOS (la lista con tema y desarrollo), redacta un texto en el que se describa detalladamente cómo se desarrolló la reunión en relación con cada tema tratado.
       - Cada tema tratado debe colocarse como subtítulo en negrilla, seguido de su respectivo desarrollo en un párrafo aparte.
       - Finalmente, incluye una conclusión general sobre los puntos abordados en la reunión, manteniendo una estructura clara y organizada, esta no debe llevar el subtitulo.
    -    {{OBJETIVO_DE_LA_REUNION_2}}: extrae el objetivo de la reunión explicado de forma clara, precisa y que no sea extensa.
    - {{COMPROMISOS_DE_REUNION}}: Esta debe ser una LISTA de objetos JSON. Cada objeto representa un compromiso de la reunion.
    - Cada objeto debe tener las claves:
        - compromiso: extrae el compromiso a realizar.
        - responsablen: extrae el nombre de la persona encargada de ejecutar el compromiso.
        - fechac: extrae la fecha de cumplimiento del compromiso.
        - fechas: extrae la fecha en la cual se va le va a hacer seguimiento al compromiso.

        Las listas deben contener objetos con las claves indicadas:
        - ASISTENTES_REUNION: nombreasistentereu, cargoasistentereunion
        - TEMAS_TRATADOS_N: tema, responsablet
        - COMPROMISOS_DE_REUNION: compromiso, resposablen, fechac, fechas
        - TEMAS_TRATADOS: tema, desarrollo
        - COMPROMISOS_R: compromiso, responsable, fechaejecucion

    TEXTO:
    ---
    {text_to_process}
    ---
    JSON:
    """
    try:
        response = model.generate_content(prompt)
        json_text = response.text.strip()
        # Limpiar posibles marcadores de código
        if json_text.startswith("```json"):
            json_text = json_text[len("```json"):].strip()
        if json_text.endswith("```"):
            json_text = json_text[:-len("```")].strip()
        match = re.search(r'\{.*\}', json_text, re.DOTALL)
        if match:
            clean_json_text = match.group(0)
            return json.loads(clean_json_text)
        else:
            st.error("⚠️ La IA no devolvió un JSON válido.")
            st.code(json_text)
            return None
    except Exception as e:
        st.error(f"Error al contactar con Gemini: {e}")
        return None

def create_word_document(template_path, data):
    try:
        doc = DocxTemplate(template_path)
        # Asegurar que las listas tengan todos los campos necesarios
        data = normalizar_listas(data)
        # Añadir campos de sesión
        data["ACTA_ELABORADA_POR"] = st.session_state.get("ACTA_ELABORADA_POR", "N/A")
        data["CARGO_ELA"] = st.session_state.get("CARGO_ELA", "N/A")
        doc.render(data)
        output_path = "acta_generada.docx"
        doc.save(output_path)
        return output_path
    except Exception as e:
        st.error(f"No se pudo generar el documento: {e}")
        return None

# ==============================================================
# INTERFAZ PRINCIPAL
# ==============================================================

logo_path = Path("logo/logo.png")
if logo_path.exists():
    with open(logo_path, "rb") as f:
        logo_base64 = base64.b64encode(f.read()).decode("utf-8")
    st.markdown(
        f"""
        <div class="app-header">
            <img src="data:image/png;base64,{logo_base64}" alt="Logo">
            <h1>Generador de Actas</h1>
        </div>
        """,
        unsafe_allow_html=True
    )
else:
    st.title("📝 Generador de Actas")

contador_actual = obtener_contador()
col1, col2, col3 = st.columns(3)
with col1:
    st.metric("Actas generadas", contador_actual)
with col2:
    st.metric("Límite", LIMITE_CONTADOR)
with col3:
    restantes = max(0, LIMITE_CONTADOR - contador_actual)
    st.metric("Restantes", restantes)

# Alerta si se alcanza el límite (comentada pero funcional)
# if contador_actual >= LIMITE_CONTADOR:
#     st.warning(f"⚠️ Se alcanzó el límite de {LIMITE_CONTADOR} actas. Es momento de reiniciar el contador.")
#     if 'alerta_enviada' not in st.session_state:
#         enviar_alerta_correo(f"Se ha alcanzado el límite de {contador_actual} actas. Debes reiniciar el contador en la app.")
#         st.session_state['alerta_enviada'] = True

if "transcripcion_area" not in st.session_state:
    st.session_state["transcripcion_area"] = ""

# Verificar existencia de plantillas
if not TEMPLATES_DIR.exists():
    st.error(f"No se encontró el directorio de plantillas: {TEMPLATES_DIR}")
    st.stop()

template_files = [f.name for f in TEMPLATES_DIR.glob("*.docx")]
if not template_files:
    st.error("No hay plantillas disponibles en la carpeta 'templates'.")
    st.stop()

template_docx = st.selectbox("📂 Selecciona una plantilla", template_files)
template_path = TEMPLATES_DIR / template_docx
template_fields = get_fields_from_template(str(template_path))

transcripcion = st.text_area("🗒️ Pega la transcripción de la reunión", height=300, key="transcripcion_area")

col1, col2 = st.columns(2)
with col1:
    st.session_state["ACTA_ELABORADA_POR"] = st.text_input("👤 Acta elaborada por")
with col2:
    st.session_state["CARGO_ELA"] = st.text_input("💼 Cargo")

col_gen, col_clear = st.columns([3, 1])
with col_gen:
    generar = st.button("📝 Generar Acta")
with col_clear:
    if st.button("🧹 Limpiar texto"):
        st.session_state["transcripcion_area"] = ""
        st.rerun()

if generar:
    if not transcripcion.strip():
        st.warning("⚠️ Debes ingresar la transcripción antes de generar.")
        st.stop()

    with st.spinner("Analizando texto con Gemini... ⏳"):
        extracted_data = extract_info_with_gemini(transcripcion, template_fields)

    if extracted_data:
        st.success("✅ Datos extraídos correctamente. Generando documento Word...")
        output_path = create_word_document(str(template_path), extracted_data)

        if output_path:
            nuevo_valor = contador_actual + 1
            actualizar_contador(nuevo_valor)
            st.success(f"🎉 Acta número {nuevo_valor} generada correctamente.")

            # if nuevo_valor >= LIMITE_CONTADOR and not st.session_state.get('alerta_enviada', False):
            #     enviar_alerta_correo(f"Se ha alcanzado el límite de {nuevo_valor} actas. Debes reiniciar el contador.")
            #     st.session_state['alerta_enviada'] = True

            with open(output_path, "rb") as f:
                st.download_button(
                    "📥 Descargar Acta Generada",
                    data=f.read(),
                    file_name=f"acta_{nuevo_valor}.docx"
                )
    else:
        st.error("No se pudo extraer información del texto.")

st.markdown("""
<div style="
    background-color: #fff0f0;
    border: 2px solid #ff9999;
    border-radius: 10px;
    padding: 15px;
    text-align: center;
    color: #660000;
    font-size: 16px;
    margin-top: 25px;
">
🚨 <b>Advertencia:</b> Esta herramienta es susceptible de mejoras. Si identifica alguna inconsistencia en el diligenciamiento del acta, por favor notifíquelo al área responsable de su diseño.
Se recomienda validar cuidadosamente toda la información generada antes de su uso, distribución o almacenamiento.<br>
</div>
""", unsafe_allow_html=True)

st.markdown("<div class='footer'>© 2025 Generador de Actas • Streamlit + Gemini + JSONBin</div>", unsafe_allow_html=True)
